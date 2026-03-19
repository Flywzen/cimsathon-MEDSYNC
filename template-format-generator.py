import sys
!{sys.executable} -m pip install python-docx
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Set margins
section = doc.sections[0]
section.top_margin = Cm(3)
section.bottom_margin = Cm(3)
section.left_margin = Cm(4)
section.right_margin = Cm(3)

# Style
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

p_format = style.paragraph_format
p_format.first_line_indent = Cm(1)
p_format.line_spacing = 1.5

# COVER
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("JUDUL KARYA (KAPITAL, TEBAL)")
run.bold = True
run.font.size = Pt(16)

doc.add_paragraph("\nCIMSAthon CAP\n").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("\n[NAMA KETUA TIM]\n[NAMA ANGGOTA]\n[NAMA ANGGOTA]").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("\n2026").alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# SECTION TEMPLATE
def add_heading(title):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

add_heading("1. Pendahuluan")
doc.add_paragraph("• Jelaskan urgensi masalah berbasis data\n• Identifikasi gap sistem\n• Kaitkan langsung ke kasus")

add_heading("2. Analisis Permasalahan (Scoring Core)")
doc.add_paragraph("• Akar masalah (root cause)\n• Fragmentasi sistem\n• Dampak klinis & sistemik\n• Evidence-based (literatur)")

add_heading("3. Konsep Solusi Digital (Scoring Terbesar)")
doc.add_paragraph("• Nama sistem\n• Arsitektur sistem\n• Alur data (input → process → output)\n• Mekanisme AI / decision support\n• Aktor (pasien, dokter, pemerintah)\n• Interoperability (FHIR/API, dll)")

add_heading("4. Implementasi di Indonesia")
doc.add_paragraph("• Infrastruktur realistis\n• Integrasi dengan BPJS / SATUSEHAT\n• Regulasi & keamanan data\n• Timeline implementasi")

add_heading("5. Dampak & Keunggulan")
doc.add_paragraph("• Outcome pasien meningkat\n• Efisiensi sistem\n• Equity & akses\n• Skalabilitas")

add_heading("6. Penutup")
doc.add_paragraph("Kesimpulan + rekomendasi kebijakan")

add_heading("Daftar Pustaka")
doc.add_paragraph("Gunakan Harvard Style")

file_path = "/content/template_CIMSAthon_scoring_optimized.docx"
doc.save(file_path)

file_path
